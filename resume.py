import streamlit as st
import pdfplumber
import docx
import re
import spacy
from docx2pdf import convert
import os

# Load spaCy model
nlp = spacy.load("en_core_web_sm")

# Predefined Skills, Job Roles, and Domains
SKILLS_LIST = {"Python", "Java", "Machine Learning", "Deep Learning", "Data Science", "SQL", "JavaScript", "AWS", "React", "TensorFlow", "Web Scraping", "Google PaLM API", "LangChain", "EDA", "Unity", "HTML", "GitHub", "C programming"}
JOB_ROLES = {"Software Engineer", "Data Scientist", "Machine Learning Engineer", "Project Manager", "AI Researcher", "Data Analyst", "Backend Developer", "Frontend Developer"}
DOMAINS = {"Healthcare", "Finance", "E-commerce", "Education", "Cybersecurity", "Robotics", "Blockchain"}

# Extract text from PDF with better layout handling
def extract_text_from_pdf(pdf_file):
    with pdfplumber.open(pdf_file) as pdf:
        text = "\n".join(page.extract_text(layout=True) for page in pdf.pages if page.extract_text(layout=True))
    return text

# Extract text from DOCX, including tables
def extract_text_from_docx(docx_file):
    doc = docx.Document(docx_file)
    text = []
    for para in doc.paragraphs:
        text.append(para.text.strip())
    
    for table in doc.tables:
        for row in table.rows:
            text.append(" ".join(cell.text.strip() for cell in row.cells))
    
    return "\n".join(text)

# Convert DOCX to PDF if needed
def convert_docx_to_pdf(docx_file):
    temp_pdf = "temp_resume.pdf"
    convert(docx_file, temp_pdf)
    return temp_pdf

# Extract Name
def extract_name(text):
    lines = text.split("\n")
    first_non_empty_line = next((line.strip() for line in lines if line.strip()), "")
    if 2 <= len(first_non_empty_line.split()) <= 4 and re.match(r'^[A-Za-z\s]+$', first_non_empty_line):
        return first_non_empty_line
    doc = nlp(text)
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            return ent.text
    return "Not Found"

# Extract Email
def extract_email(text):
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    match = re.search(email_pattern, text)
    return match.group(0) if match else "Not Found"

# Extract Phone Number
def extract_phone(text):
    phone_pattern = r'\b(?:\+?\d{1,3})?\s?(?:\(?\d{3}\)?[\s.-]?)?\d{3}[\s.-]?\d{4}\b'
    match = re.search(phone_pattern, text)
    return match.group(0) if match else "Not Found"

# Extract Skills
def extract_skills(text):
    found_skills = {skill for skill in SKILLS_LIST if skill.lower() in text.lower()}
    return ", ".join(sorted(found_skills)) if found_skills else "Not Found"

# Extract Experience
import re
from datetime import datetime

def extract_experience(text):
    current_year = datetime.now().year
    total_experience = 0
    seen_periods = set()

    # **Enhanced Regex Patterns for Work Experience**
    experience_patterns = [
        r'(?i)(\b\d{4}\b)\s*[-–]\s*(\b\d{4}|\bPresent\b|\bCurrent\b)',  # YYYY - YYYY / YYYY - Present
        r'(?i)(\b\d{4}\b)\s+to\s+(\b\d{4}|\bPresent\b|\bCurrent\b)',  # YYYY to YYYY / Present
        r'(?i)([A-Za-z]+\s+\d{4})\s*[-–]\s*([A-Za-z]+\s+\d{4}|\bPresent\b|\bCurrent\b)',  # Month YYYY - Month YYYY / Present
    ]

    # **Strict Exclusion Keywords**
    education_keywords = ["bachelor", "master", "phd", "diploma", "college", "university", "education", "cgpa"]
    internship_keywords = ["intern", "internship", "virtual internship"]
    project_keywords = ["project", "capstone", "research", "thesis", "hackathon", "simulation"]

    # **Extract Education Years**
    education_patterns = [
        r'(?i)(bachelor|master|phd|diploma|college|university|education).*?(\d{4})\s*[-–]\s*(\d{4})',
        r'(?i)(cgpa|gpa|percentage|ssc|hsc|class\s?(x|xii)).*?(\d{4})\s*[-–]\s*(\d{4})'
    ]
    
    education_years = set()
    for pattern in education_patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            start_year, end_year = int(match[-2]), int(match[-1])
            education_years.update(range(start_year, end_year + 1))

    #print("Extracted education years:", education_years)

    # **Ensure "Experience" Section Exists**
    if not re.search(r'(?i)\b(work\s+experience|professional\s+experience|employment\s+history)\b', text):
        print("⚠️ No experience section found. Returning 0 years.")
        return "0 years"

    #print("✅ Experience section detected!")

    # **Extract Work Experience**
    for pattern in experience_patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            try:
                start_year, end_year = match
                start_year = int(re.search(r'\d{4}', start_year).group())
                end_year = current_year if re.search(r'(?i)present|current', end_year) else int(re.search(r'\d{4}', end_year).group())

                #print(f"🔍 Detected Experience: Start={start_year}, End={end_year}")

                # **Ignore education years**
                if start_year in education_years or end_year in education_years:
                    #print(f"❌ Ignoring {start_year}-{end_year} (Education detected)")
                    continue

                # **Check If This Is an Internship (Only Ignore If Near "Intern")**
                experience_text_snippet = text.lower()
                snippet_start = max(text.lower().find(str(start_year)) - 50, 0)
                snippet_end = min(text.lower().find(str(end_year)) + 50, len(text))
                context_snippet = text.lower()[snippet_start:snippet_end]

                if any(keyword in context_snippet for keyword in internship_keywords):
                    #print(f"❌ Ignoring {start_year}-{end_year} (Internship detected in job title)")
                    continue

                # **Allow Multiple Work Experiences (Fix Duplicate Issue)**
                if (start_year, end_year) in seen_periods:
                    #print(f"⚠️ Skipping {start_year}-{end_year} (Duplicate detected, but allowing other entries)")
                else:
                    seen_periods.add((start_year, end_year))

                    # **Calculate Experience Duration**
                    years_of_experience = max(0, end_year - start_year)
                    total_experience += years_of_experience
                    #print(f"✅ Added {years_of_experience} years. Total so far: {total_experience}")

            except ValueError:
                #print(f"⚠️ Skipping invalid date range: {match}")
                continue

    #print(f"🎯 Final total experience: {total_experience} years")
    return f"{total_experience} years" if total_experience > 0 else "0 years"


# Extract Job Role
def extract_job_role(text):
    doc = nlp(text)
    found_roles = set()
    for token in doc:
        if token.text in JOB_ROLES:
            found_roles.add(token.text)
    
    about_me_patterns = ["Software Engineer", "Data Scientist", "Machine Learning Engineer", "AI Engineer", 
    "Deep Learning Engineer", "Data Analyst", "Business Analyst", "Backend Developer", 
    "Frontend Developer", "Full Stack Developer", "Cloud Engineer", "DevOps Engineer", 
    "Cybersecurity Analyst", "Embedded Systems Engineer", "Blockchain Developer", 
    "Computer Vision Engineer", "NLP Engineer", "Product Manager", "QA Engineer", 
    "Automation Engineer", "Data Engineer", "Database Administrator", "Software Architect", 
    "System Administrator", "IT Support Engineer", "Game Developer", "Mobile App Developer", 
    "AR/VR Developer", "Site Reliability Engineer", "AI Research Scientist", "Big Data Engineer", 
    "BI Analyst", "Security Engineer", "Network Engineer", "IoT Engineer", 
    "Technical Support Engineer", "Research Scientist", "Technical Program Manager", 
    "DataOps Engineer", "MLOps Engineer", "Robotics Engineer", "Quantum Computing Engineer",]
    for phrase in about_me_patterns:
        if phrase in text.lower():
            found_roles.add(phrase.title())
    
    return ", ".join(found_roles) if found_roles else "Not Specified"

# Extract Domain
def extract_domain(text):
    found_domains = {domain for domain in DOMAINS if domain.lower() in text.lower()}
    return ", ".join(sorted(found_domains)) if found_domains else "Not Found"

# Streamlit App
st.title("Resume Parser")

uploaded_file = st.file_uploader("Upload Resume (PDF or DOCX)", type=["pdf", "docx"])

if uploaded_file:
    extracted_text = ""
    
    if uploaded_file.type == "application/pdf":
        extracted_text = extract_text_from_pdf(uploaded_file)
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        temp_pdf = convert_docx_to_pdf(uploaded_file)
        extracted_text = extract_text_from_pdf(temp_pdf)
        os.remove(temp_pdf)  # Clean up temp file
    else:
        st.error("Unsupported file type")
        st.stop()
    
    # Extract Information
    name = extract_name(extracted_text)
    email = extract_email(extracted_text)
    phone = extract_phone(extracted_text)
    skills = extract_skills(extracted_text)
    experience = extract_experience(extracted_text)
    job_role = extract_job_role(extracted_text)
    domain = extract_domain(extracted_text)
    
    # Display Results
    st.subheader("Extracted Information")
    st.write(f"**Name:** {name}")
    st.write(f"**Email:** {email}")
    st.write(f"**Phone:** {phone}")
    st.write(f"**Skills:** {skills}")
    st.write(f"**Experience:** {experience}")
    st.write(f"**Job Role:** {job_role}")
    st.write(f"**Domain:** {domain}")
